# ai_services/views.py
"""
AI Services views — expose AI analysis, resume building, and cover letter endpoints.
"""
from django.shortcuts import render, get_object_or_404, redirect
from django.http import JsonResponse, HttpResponse
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.views.decorators.http import require_POST

from resumes.models import Resume
from jobs.models import Job
from applications.models import Application

from .gemini_service import GeminiService
from .grok_service   import GrokService
from .fusion_service import FusionEngine
from .ats_engine     import ATSEngine
from .matcher        import ResumeJobMatcher


def _run_full_analysis(resume) -> None:
    """
    Run the complete AI pipeline on a Resume object and save all results.
    Called after upload or on-demand re-analysis.
    """
    # Ensure raw text is extracted from file if it was saved empty
    if not resume.raw_text or not resume.raw_text.strip():
        try:
            from resumes.services import ResumeParser
            parser = ResumeParser()
            parsed_data = parser.parse(resume.file.path)
            resume.raw_text     = parsed_data.get('raw_text', '')
            resume.skills       = parsed_data.get('skills', [])
            resume.experience   = parsed_data.get('experience', [])
            resume.education    = parsed_data.get('education', [])
            resume.contact_info = parsed_data.get('contact_info', {})
            resume.save()
        except Exception as e:
            print(f"[Analysis Setup] Failed to parse PDF: {e}")

    resume_text = resume.get_resume_text()

    gemini  = GeminiService()
    grok    = GrokService()
    ats_eng = ATSEngine()
    fusion  = FusionEngine()

    # 1. Parallel analyses
    gemini_result = gemini.analyze_resume(resume_text)
    grok_result   = grok.analyze_technical(resume_text)
    ats_result    = ats_eng.score(resume_text)

    # 2. Fuse
    fused = fusion.fuse_resume_analysis(gemini_result, grok_result)

    # 3. Override ATS score with the rule-based engine (more reliable)
    fused['ats_score'] = ats_result.get('ats_score', fused.get('ats_score', 0))

    # 4. Save all to model
    resume.gemini_analysis = gemini_result
    resume.grok_analysis   = grok_result
    resume.fusion_analysis = fused
    resume.ai_suggestions  = fused           # backwards compat

    resume.resume_score    = fused.get('resume_score', 0)
    resume.ats_score       = fused.get('ats_score', 0)
    resume.grammar_score   = fused.get('grammar_score', 0)
    resume.technical_score = fused.get('technical_score', 0)

    resume.save()


# ─────────────────────────────────────────────────────────────────
# EXISTING ENDPOINTS (updated)
# ─────────────────────────────────────────────────────────────────

@login_required
def analyze_resume(request, resume_id):
    """Re-run full AI analysis pipeline on a resume."""
    resume = get_object_or_404(Resume, id=resume_id, user=request.user)
    try:
        _run_full_analysis(resume)
        messages.success(request, '✅ AI analysis complete — all scores updated.')
    except Exception as e:
        messages.warning(request, f'Analysis partially complete: {e}')
    return redirect('resume_detail', resume_id=resume.id)


@login_required
def match_job(request, job_id):
    """Match user's primary resume to a job (JSON API)."""
    job    = get_object_or_404(Job, id=job_id)
    resume = Resume.objects.filter(user=request.user).order_by('-created_at').first()
    if not resume:
        return JsonResponse({'error': 'Please upload a resume first'}, status=400)

    matcher = ResumeJobMatcher()
    result  = matcher.match(resume, job)
    return JsonResponse({'success': True, 'match': result})


@login_required
def rank_candidates(request, job_id):
    """Rank all applicants for a job (JSON API — Recruiter only)."""
    if request.user.user_type != 'recruiter':
        return JsonResponse({'error': 'Only recruiters can rank candidates'}, status=403)

    job          = get_object_or_404(Job, id=job_id, recruiter=request.user)
    applications = Application.objects.filter(job=job).select_related('candidate', 'resume')
    matcher      = ResumeJobMatcher()
    ranked       = matcher.rank_candidates(job, applications)

    result = []
    for item in ranked:
        app = item['application']
        try:
            cname = f"{app.candidate.candidate_profile.first_name} {app.candidate.candidate_profile.last_name}"
        except Exception:
            cname = app.candidate.username
        result.append({
            'rank'               : item['rank'],
            'candidate'          : cname,
            'email'              : app.candidate.email,
            'score'              : item['score'],
            'hire_probability'   : item['hire_probability'],
            'recommendation'     : item['recommendation'],
            'matched_skills'     : item['matched_skills'],
            'missing_skills'     : item['missing_skills'],
            'interview_questions': item['interview_questions'],
        })

    return JsonResponse({'success': True, 'candidates': result})


@login_required
def generate_voice(request, resume_id):
    """Return voice text for browser Speech API (no server-side TTS)."""
    resume   = get_object_or_404(Resume, id=resume_id, user=request.user)
    analysis = resume.fusion_analysis or resume.ai_suggestions or {}

    score    = analysis.get('resume_score', resume.resume_score or 0)
    strengths = '; '.join((analysis.get('strengths') or [])[:2])
    suggest  = '; '.join((analysis.get('suggestions') or [])[:2])

    text = (
        f"Hello! Here is your resume analysis. "
        f"Your resume score is {score} out of 100. "
        f"Your strengths include: {strengths}. "
        f"Suggestions for improvement: {suggest}. "
        f"Good luck with your job search!"
    )
    return JsonResponse({'success': True, 'text': text, 'audio_url': None})


@login_required
def generate_match_voice(request, job_id):
    """Return voice text for a job match result."""
    job         = get_object_or_404(Job, id=job_id)
    application = Application.objects.filter(candidate=request.user, job=job).first()
    if not application:
        return JsonResponse({'error': 'No application found'}, status=404)

    m = application.fusion_match or application.match_breakdown or {}
    text = (
        f"Job Match Analysis. "
        f"Your resume matches {job.title} at {job.company} by {m.get('overall_score', 0)} percent. "
        f"Matched skills: {', '.join((m.get('matched_skills') or [])[:4])}. "
        f"Skills to add: {', '.join((m.get('missing_skills') or [])[:3])}. "
        f"{m.get('recommendation', '')}."
    )
    return JsonResponse({'success': True, 'text': text, 'audio_url': None})


# ─────────────────────────────────────────────────────────────────
# NEW v2.0 ENDPOINTS
# ─────────────────────────────────────────────────────────────────

@login_required
@require_POST
def build_resume(request, resume_id):
    """Generate a professional ATS-optimised resume text via Gemini + Grok."""
    resume      = get_object_or_404(Resume, id=resume_id, user=request.user)
    target_role = request.POST.get('target_role', '')

    gemini = GeminiService()
    grok   = GrokService()
    
    # Step 1: Gemini creates initial professional draft
    draft  = gemini.build_resume(resume.get_resume_text(), target_role)
    
    # Step 2: Grok audits the draft and recommends better technical wording/tools
    suggestions = grok.suggest_technologies(draft, target_role)
    
    # Step 3: Gemini synthesizes suggestions to build final optimized resume
    final_resume = gemini.synthesize_final_resume(draft, target_role, suggestions)

    resume.built_resume = final_resume
    resume.save(update_fields=['built_resume'])

    messages.success(request, '✅ Professional resume generated by Gemini & Grok!')
    return redirect('resume_detail', resume_id=resume.id)


@login_required
def download_built_resume(request, resume_id):
    """Download the AI-built resume as a .txt file."""
    resume = get_object_or_404(Resume, id=resume_id, user=request.user)
    if not resume.built_resume:
        messages.warning(request, 'Please generate the resume first.')
        return redirect('resume_detail', resume_id=resume.id)

    response = HttpResponse(resume.built_resume, content_type='text/plain')
    filename = f"resume_{resume.id}_professional.txt"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@login_required
def download_built_resume_docx(request, resume_id):
    """Download the AI-built resume as a styled Microsoft Word .docx file."""
    resume = get_object_or_404(Resume, id=resume_id, user=request.user)
    if not resume.built_resume:
        messages.warning(request, 'Please generate the resume first.')
        return redirect('resume_detail', resume_id=resume.id)

    from docx import Document
    from io import BytesIO

    doc = Document()
    
    # ── Style adjustments ──
    # Document Title / Name
    title_p = doc.add_paragraph()
    run = title_p.add_run(resume.title)
    run.font.size = 24 * 12700  # Pt size (roughly) or use docx.shared.Pt if preferred
    run.bold = True
    title_p.alignment = 1  # Centered
    
    lines = resume.built_resume.split('\n')
    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue
            
        # Check if this line looks like a header
        if (stripped.isupper() and len(stripped) < 40) or (stripped.startswith('###') or stripped.startswith('##')):
            clean_hdr = stripped.replace('#', '').strip()
            hdr = doc.add_heading(clean_hdr, level=2)
            hdr.paragraph_format.space_before = 180000  # Space before
            hdr.paragraph_format.space_after = 60000    # Space after
        elif stripped.startswith('==') or stripped.startswith('--'):
            continue  # ignore plain dividers
        elif stripped.startswith('-') or stripped.startswith('*'):
            clean_item = stripped[1:].strip()
            doc.add_paragraph(clean_item, style='List Bullet')
        else:
            doc.add_paragraph(stripped)

    # Save to a memory stream
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    filename = f"resume_{resume.id}_professional.docx"
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response



@login_required
@require_POST
def generate_cover_letter(request, resume_id):
    """Generate a cover letter for a specific job via Gemini."""
    resume  = get_object_or_404(Resume, id=resume_id, user=request.user)
    job_id  = request.POST.get('job_id')
    job     = get_object_or_404(Job, id=job_id) if job_id else None

    gemini = GeminiService()
    letter = gemini.generate_cover_letter(
        resume_text  = resume.get_resume_text(),
        job_title    = job.title    if job else request.POST.get('job_title', 'Software Developer'),
        company      = job.company  if job else request.POST.get('company', 'the company'),
        requirements = job.requirements if job else '',
    )

    resume.cover_letter = letter
    resume.save(update_fields=['cover_letter'])

    messages.success(request, '✅ Cover letter generated!')
    return redirect('resume_detail', resume_id=resume.id)


@login_required
def download_cover_letter(request, resume_id):
    """Download the generated cover letter as a .txt file."""
    resume = get_object_or_404(Resume, id=resume_id, user=request.user)
    if not resume.cover_letter:
        messages.warning(request, 'Please generate a cover letter first.')
        return redirect('resume_detail', resume_id=resume.id)

    response = HttpResponse(resume.cover_letter, content_type='text/plain')
    response['Content-Disposition'] = f'attachment; filename="cover_letter_{resume.id}.txt"'
    return response


@login_required
def interview_questions(request, app_id):
    """Return interview questions for an application (Recruiter view)."""
    if request.user.user_type != 'recruiter':
        return JsonResponse({'error': 'Recruiters only'}, status=403)

    app = get_object_or_404(Application, id=app_id)
    if app.job.recruiter != request.user:
        return JsonResponse({'error': 'Access denied'}, status=403)

    questions = app.interview_questions or []

    # Generate if not present
    if not questions and app.resume:
        grok   = GrokService()
        result = grok.rank_candidate(app.resume.get_resume_text(), app.job, app.match_score)
        questions = result.get('interview_questions', [])
        app.interview_questions = questions
        app.hire_probability    = result.get('hire_probability', 0)
        app.save(update_fields=['interview_questions', 'hire_probability'])

    return JsonResponse({'success': True, 'questions': questions, 'application_id': app_id})


@login_required
def analytics_dashboard(request):
    """Render or return recruitment analytics visualization payload."""
    from .analytics import RecruitmentAnalyticsEngine
    engine = RecruitmentAnalyticsEngine()
    data = engine.get_dashboard_data()

    if request.GET.get('format') == 'json' or request.headers.get('x-requested-with') == 'XMLHttpRequest':
        return JsonResponse(data)

    return render(request, 'analytics/dashboard.html', {'analytics': data})